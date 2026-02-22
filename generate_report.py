"""
Generate DOCX report for OZ Different - Period Poverty Research in Bardejov
"""

import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import seaborn as sns
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# ─── Paths ───
BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE, '_report_images')
os.makedirs(IMG_DIR, exist_ok=True)

# ─── Load data ───
pre_data = pd.read_csv(os.path.join(BASE, 'pre_installation_data.csv'))
pre_data = pre_data.map(lambda x: x.strip() if isinstance(x, str) else x)

after_data = pd.read_csv(os.path.join(BASE, 'after_installation_data.csv'))
after_data = after_data.map(lambda x: x.strip() if isinstance(x, str) else x)

# Pre-data column renaming (same as notebook)
pre_data = pre_data.drop(columns=[
    'Kde alebo od koho ste získali informácie o menštruácii? (môžete zaškrtnúť viac možností)',
    'Aké menštruačné pomôcky ste používali? (môžete zaškrtnúť viac možností)',
    'S akými prekážkami ste sa počas menštruácie najčastejšie stretli?',
    'Aké pocity alebo emócie najčastejšie pociťujete počas menštruácie? (napíšte):',
    'Ak máte podozrenie na gynekologický problém, kde najskôr hľadáte informácie? (napíšte)',
    'Priestor na Vaše pripomienky a komentáre (NEPOVINNÉ):'
])

pre_data.columns = [
    'Timestamp','Vek', 'Škola', 'Ročník', 'S kým aktuálne bývate?', 'Rodinný stav', 'Počet detí', 'Počet bratov', 'Počet sestier', 'Počet súrodencov',
    'Zamestnanie otca', 'Najvyššie dosiahnuté vzdelanie alebo posledný ukončený ročník otca', 'Zamestnanie matky', 'Najvyššie dosiahnuté vzdelanie alebo posledný ukončený ročník matky',
    'Prístup k teplej vode', 'Prístup k sprche alebo vani', 'Prístup k splachovaciemu WC', 'Prístup ku teplu alebo kúreniu',
    'Mávate aktuálne menštruáciu', 'Vek prvej menštruácie', 'Mali ste pred prvou menštruáciou dostatok informácií o tom, čo menštruácia znamená a ako sa na ňu pripraviť?',
    'Informácie o menštruácií získané od iného rodinného príslušníka', 'Informácie o menštruácií získané zo školy', 'Informácie o menštruácií získané od sestry/sestier', 'Informácie o menštruácií získané z prednášok/workshopov', 'Informácie o menštruácií získané od kamarátov', 'Informácie o menštruácií získané z internetu', 'Informácie o menštruácií získané od matky',
    'Používané potreby: Handry','Používané potreby: Menštruačné nohavičky','Používané porteby: Intímky','Používané porteby: Tampóny','Používané potreby: Menštruačné vložky',
    'Dostatok pomôcok na celé trvanie menštruácie', 'Prekážka: peniaze', 'Prekážka: žiadne', 'Prekážka: bolesť',
    'Sledujete svoj menštruačný cyklus?','Akým spôsobom si zaznamenávate svoj cyklus?', 'Vnímate menštruáciu ako zásah do svojich každodenných plánov?',
    'Pocity: smútok / depresia / úzkosť / strach', 'Pocity: hnev / nervozita / náladovosť / stres', 'Pocity: únava', 'Pocity: bolesť',
    'Informácie ku gynekologickému problému získané z/od : Lekára', 'Informácie ku gynekologickému problému získané z/od : Kamarátov', 'Informácie ku gynekologickému problému získané z/od : Internetu', 'Informácie ku gynekologickému problému získané z/od : Mamy',
    'Je pre vás ťažké komunikovať o intímnych témach so svojím lekárom?', 'Pri hľadaní informácií o zdravotných problémoch dávate prednosť:', 'Nosievate so sebou zásobu menštruačných pomôcok ako prvú pomoc?',
    'Je pre vás výmena vložky alebo tampónu stresujúca, ak ste mimo domova?', 'Cítili ste sa niekedy trápne pri nákupe menštruačných pomôcok?', 'Stalo sa vám, že ste si kvôli finančným dôvodom nemohli dovoliť kúpiť menštruačné pomôcky?',
    'Vynechali ste niekedy školu kvôli menštruácii?', 'Ako vnímate menštruáciu?'
]

# Setup
yes_no_map = {
    'Áno': 1, 'Yes': 1,
    'Nie': 0, 'No': 0,
    'Niekedy': 0.5, 'Sometimes': 0.5,
    'Nechcem odpovedať': np.nan, "Don't want to answer": np.nan
}

access_cols = ['Prístup k teplej vode', 'Prístup k sprche alebo vani',
               'Prístup k splachovaciemu WC', 'Prístup ku teplu alebo kúreniu']
pre_data['Lack_count'] = pre_data[access_cols].apply(lambda row: (row == 'Nie').sum(), axis=1)

answer_map_sk = {
    'Áno': 'Áno',
    'Nie': 'Nie',
    'Nechcem odpovedať': 'Nechcem odpovedať'
}

info_prep_map = {
    'Áno, mala som všetky potrebné informácie': 'Áno, mala som všetky potrebné informácie',
    'Mala som len čiastočné informácie': 'Mala som len čiastočné informácie',
    'Nemala som žiadne informácie': 'Nemala som žiadne informácie'
}

# ─── Chart generation helpers ───
CHART_COLOR = '#1a4a6e'
CHART_COLOR2 = '#6baed6'
COLORS_COMPARISON = ['#2171b5', '#6baed6']

def save_fig(name):
    path = os.path.join(IMG_DIR, f'{name}.png')
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


# ═══════════════════════════════════════════
# GENERATE ALL CHARTS
# ═══════════════════════════════════════════

num_pre = len(pre_data)
num_after = len(after_data)
avg_age = pre_data['Vek'].mean().__round__(2)
avg_first_period_age = pre_data['Vek prvej menštruácie'].mean().__round__(2)

# --- PRE 1: Age distribution ---
plt.figure(figsize=(10, 6))
plt.hist(pre_data['Vek'], bins=range(12, 21), edgecolor='black', alpha=0.9, color=CHART_COLOR)
plt.axvline(x=avg_age, color='#fffacd', linestyle='--', linewidth=2, label=f'Priemer: {avg_age:.2f}')
plt.xlabel('Vek')
plt.ylabel('Počet respondentiek')
plt.title('Rozdelenie veku respondentiek')
plt.legend()
plt.xticks([x + 0.5 for x in range(12, 20)], range(12, 20))
plt.tight_layout()
img_pre_age = save_fig('pre_age')

# --- PRE 2: Age of first period ---
plt.figure(figsize=(10, 6))
plt.hist(pre_data['Vek prvej menštruácie'], bins=range(8, 17), edgecolor='black', alpha=0.9, color=CHART_COLOR)
plt.axvline(x=avg_first_period_age, color='#fffacd', linestyle='--', linewidth=2, label=f'Priemer: {avg_first_period_age:.2f}')
plt.xlabel('Vek prvej menštruácie')
plt.ylabel('Počet respondentiek')
plt.title('Rozdelenie veku prvej menštruácie')
plt.legend()
plt.xticks([x + 0.5 for x in range(9, 16)], range(9, 16))
plt.tight_layout()
img_pre_first_period = save_fig('pre_first_period')

# --- PRE 3: Missed school ---
missed_counts = pre_data['Vynechali ste niekedy školu kvôli menštruácii?'].map(answer_map_sk).value_counts()
order = ['Nechcem odpovedať', 'Nie', 'Áno']
missed_counts = missed_counts.reindex([x for x in order if x in missed_counts.index])

fig, ax = plt.subplots(figsize=(10, 4))
bars = ax.barh(missed_counts.index, missed_counts.values, color=CHART_COLOR)
total = num_pre
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{v}}}$ ({v/total*100:.1f}%)' for v in missed_counts.values])
ax.xaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Vynechanie školy kvôli menštruácii')
plt.tight_layout()
img_pre_missed = save_fig('pre_missed_school')

# --- PRE 4: Affordability ---
afford_counts = pre_data['Stalo sa vám, že ste si kvôli finančným dôvodom nemohli dovoliť kúpiť menštruačné pomôcky?'].map(answer_map_sk).value_counts()
fig, ax = plt.subplots(figsize=(10, 4))
bars = ax.barh(afford_counts.index, afford_counts.values, color=CHART_COLOR)
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{v}}}$ ({v/total*100:.1f}%)' for v in afford_counts.values])
ax.xaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Nemožnosť kúpiť si menštruačné pomôcky z finančných dôvodov aspoň raz')
plt.tight_layout()
img_pre_afford = save_fig('pre_afford')

# --- PRE 5: Information preparedness ---
info_prep_counts = pre_data['Mali ste pred prvou menštruáciou dostatok informácií o tom, čo menštruácia znamená a ako sa na ňu pripraviť?'].map(info_prep_map).value_counts()
fig, ax = plt.subplots(figsize=(10, 4))
bars = ax.barh(info_prep_counts.index, info_prep_counts.values, color=CHART_COLOR)
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{v}}}$ ({v/total*100:.1f}%)' for v in info_prep_counts.values])
ax.xaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Dostatok informácií pred prvou menštruáciou')
plt.tight_layout()
img_pre_info_prep = save_fig('pre_info_prep')

# --- PRE 6: Information sources ---
info_cols = {
    'Informácie o menštruácií získané od matky': 'Mama',
    'Informácie o menštruácií získané zo školy': 'Škola',
    'Informácie o menštruácií získané z internetu': 'Internet',
    'Informácie o menštruácií získané od kamarátov': 'Kamarátky',
    'Informácie o menštruácií získané od sestry/sestier': 'Sestra/sestry',
    'Informácie o menštruácií získané od iného rodinného príslušníka': 'Iný rodinný príslušník',
    'Informácie o menštruácií získané z prednášok/workshopov': 'Prednášky/Workshopy'
}
info_sums = pre_data[list(info_cols.keys())].sum().sort_values(ascending=True)
info_sums.index = [info_cols[col] for col in info_sums.index]

fig, ax = plt.subplots(figsize=(10, 5))
bars = ax.barh(info_sums.index, info_sums.values, color=CHART_COLOR)
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{int(v)}}}$ ({v/num_pre*100:.1f}%)' for v in info_sums.values])
ax.xaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Zdroje informácií o menštruácii')
plt.tight_layout()
img_pre_info_sources = save_fig('pre_info_sources')

# --- PRE 7: Info preparedness vs age of first period ---
df_analysis = pre_data[['Mali ste pred prvou menštruáciou dostatok informácií o tom, čo menštruácia znamená a ako sa na ňu pripraviť?', 'Vek prvej menštruácie']].copy()
df_analysis.columns = ['Úroveň informovanosti', 'Vek prvej menštruácie']
df_analysis['Úroveň informovanosti'] = df_analysis['Úroveň informovanosti'].map(info_prep_map)
mean_ages = df_analysis.groupby('Úroveň informovanosti')['Vek prvej menštruácie'].mean()

fig, ax = plt.subplots(figsize=(10, 5))
bars = ax.bar(mean_ages.index, mean_ages.values, color=CHART_COLOR)
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{v:.1f}}}$ rokov' for v in mean_ages.values])
ax.yaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Priemerný vek prvej menštruácie podľa úrovne informovanosti')
ax.set_xlabel('Úroveň informovanosti pred prvou menštruáciou')
plt.tight_layout()
img_pre_info_age = save_fig('pre_info_age')

# --- PRE 8: Products used ---
product_cols = {
    'Používané potreby: Menštruačné vložky': 'Menštruačné vložky',
    'Používané porteby: Tampóny': 'Tampóny',
    'Používané potreby: Menštruačné nohavičky': 'Menštruačné nohavičky',
    'Používané porteby: Intímky': 'Intímky',
    'Používané potreby: Handry': 'Handry'
}
product_sums = pre_data[list(product_cols.keys())].sum().sort_values(ascending=True)
product_sums.index = [product_cols[col] for col in product_sums.index]

fig, ax = plt.subplots(figsize=(10, 5))
bars = ax.barh(product_sums.index, product_sums.values, color=CHART_COLOR)
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{int(v)}}}$ ({v/num_pre*100:.1f}%)' for v in product_sums.values])
ax.xaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Používané menštruačné pomôcky')
plt.tight_layout()
img_pre_products = save_fig('pre_products')

# --- PRE 9: Access to amenities ---
columns_amenities = {
    'Prístup k teplej vode': 'Prístup k teplej vode',
    'Prístup k sprche alebo vani': 'Prístup k sprche alebo vani',
    'Prístup k splachovaciemu WC': 'Prístup k splachovaciemu WC',
    'Prístup ku teplu alebo kúreniu': 'Prístup ku kúreniu'
}

data_amenities = {}
for sk_col, label in columns_amenities.items():
    counts = pre_data[sk_col].map(answer_map_sk).value_counts()
    data_amenities[label] = counts

df_plot = pd.DataFrame(data_amenities).T
df_plot = df_plot.reindex(columns=['Áno', 'Nie', 'Nechcem odpovedať']).fillna(0)

full_access = (pre_data['Lack_count'] == 0).sum()
lacking_any = (pre_data['Lack_count'] > 0).sum()

fig, ax = plt.subplots(figsize=(10, 6))
y = np.arange(len(df_plot))
height = 0.25
bars1 = ax.barh(y + height, df_plot['Áno'], height, label='Áno', color='#6baed6')
bars2 = ax.barh(y, df_plot['Nie'], height, label='Nie', color='#2171b5')
bars3 = ax.barh(y - height, df_plot['Nechcem odpovedať'], height, label='Nechcem odpovedať', color='#08306b')

ax.bar_label(bars1, padding=3, labels=[f'{v:.0f} ({v/num_pre*100:.1f}%)' if v > 0 else '' for v in df_plot['Áno']])
ax.bar_label(bars2, padding=3, labels=[f'{v:.0f} ({v/num_pre*100:.1f}%)' if v > 0 else '' for v in df_plot['Nie']])
ax.bar_label(bars3, padding=3, labels=[f'{v:.0f} ({v/num_pre*100:.1f}%)' if v > 0 else '' for v in df_plot['Nechcem odpovedať']])
ax.text(0.95, 0.05, f'Plný prístup: {full_access} ({full_access/num_pre*100:.1f}%)\nChýba ≥1: {lacking_any} ({lacking_any/num_pre*100:.1f}%)',
        transform=ax.transAxes, ha='right', va='bottom', fontsize=10,
        bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))
ax.xaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Prístup k vybavenosti')
ax.set_yticks(y)
ax.set_yticklabels(df_plot.index)
ax.legend()
plt.tight_layout()
img_pre_amenities = save_fig('pre_amenities')

# --- PRE 10: Amenities by siblings ---
def sibling_group(n):
    if pd.isna(n): return None
    elif n == 0: return '0'
    elif n <= 2: return '1-2'
    elif n <= 4: return '3-4'
    else: return '5+'

pre_data['Sibling_group'] = pre_data['Počet súrodencov'].apply(sibling_group)
group_order = ['0', '1-2', '3-4', '5+']
plot_data = pre_data[pre_data['Sibling_group'].notna()]
group_means = plot_data.groupby('Sibling_group')['Lack_count'].mean()
group_counts = plot_data.groupby('Sibling_group')['Lack_count'].count()

fig, ax = plt.subplots(figsize=(10, 5))
bars = ax.bar([g for g in group_order if g in group_means.index],
              [group_means[g] for g in group_order if g in group_means.index], color=CHART_COLOR)
for i, g in enumerate([g for g in group_order if g in group_means.index]):
    ax.text(i, group_means[g] + 0.05, f'$\\mathbf{{{group_means[g]:.2f}}}$\n(n={group_counts[g]})',
            ha='center', va='bottom', fontsize=10)
ax.set_xlabel('Počet súrodencov')
ax.set_title('Priemerný počet chýbajúcich vybaveností podľa počtu súrodencov')
ax.yaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
plt.tight_layout()
img_pre_siblings = save_fig('pre_siblings_amenities')

# --- PRE 11: Amenities by age ---
def age_group(n):
    if pd.isna(n): return None
    elif n <= 13: return '12-13'
    elif n <= 15: return '14-15'
    elif n <= 17: return '16-17'
    else: return '18-19'

pre_data['Age_group'] = pre_data['Vek'].apply(age_group)
group_order_age = ['12-13', '14-15', '16-17', '18-19']
plot_data_age = pre_data[pre_data['Age_group'].notna()]
group_means_age = plot_data_age.groupby('Age_group')['Lack_count'].mean()
group_counts_age = plot_data_age.groupby('Age_group')['Lack_count'].count()

fig, ax = plt.subplots(figsize=(10, 5))
bars = ax.bar([g for g in group_order_age if g in group_means_age.index],
              [group_means_age[g] for g in group_order_age if g in group_means_age.index], color=CHART_COLOR)
for i, g in enumerate([g for g in group_order_age if g in group_means_age.index]):
    ax.text(i, group_means_age[g] + 0.05, f'$\\mathbf{{{group_means_age[g]:.2f}}}$\n(n={group_counts_age[g]})',
            ha='center', va='bottom', fontsize=10)
ax.set_xlabel('Veková skupina')
ax.set_title('Priemerný počet chýbajúcich vybaveností podľa vekovej skupiny')
ax.yaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
plt.tight_layout()
img_pre_age_amenities = save_fig('pre_age_amenities')

# --- PRE 12: Symptoms ---
symptom_cols = {
    'Pocity: bolesť': 'Bolesť',
    'Pocity: únava': 'Únava',
    'Pocity: hnev / nervozita / náladovosť / stres': 'Hnev / Nervozita / Náladovosť / Stres',
    'Pocity: smútok / depresia / úzkosť / strach': 'Smútok / Depresia / Úzkosť / Strach'
}
symptom_sums = pre_data[list(symptom_cols.keys())].sum().sort_values(ascending=True)
symptom_sums.index = [symptom_cols[col] for col in symptom_sums.index]

fig, ax = plt.subplots(figsize=(10, 5))
bars = ax.barh(symptom_sums.index, symptom_sums.values, color=CHART_COLOR)
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{int(v)}}}$ ({v/num_pre*100:.1f}%)' for v in symptom_sums.values])
ax.xaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Symptómy pociťované počas menštruácie')
plt.tight_layout()
img_pre_symptoms = save_fig('pre_symptoms')

# --- PRE 13: Tampon users hot water ---
tampon_users = pre_data[pre_data['Používané porteby: Tampóny'] == 1]
hot_water_counts = tampon_users['Prístup k teplej vode'].map(answer_map_sk).value_counts()
order_hw = ['Nechcem odpovedať', 'Nie', 'Áno']
hot_water_counts = hot_water_counts.reindex([x for x in order_hw if x in hot_water_counts.index])

fig, ax = plt.subplots(figsize=(10, 4))
bars = ax.barh(hot_water_counts.index, hot_water_counts.values, color=CHART_COLOR)
total_tampon = len(tampon_users)
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{v}}}$ ({v/total_tampon*100:.1f}%)' for v in hot_water_counts.values])
ax.xaxis.set_visible(False)
for spine in ax.spines.values():
    spine.set_visible(False)
ax.set_title('Prístup k teplej vode medzi používateľkami tampónov')
plt.tight_layout()
img_pre_tampon_water = save_fig('pre_tampon_water')


# ═══════════════════════════════════════════
# AFTER INSTALLATION CHARTS
# ═══════════════════════════════════════════

answer_map_after = {
    'Ano': 'Áno',
    'Nie': 'Nie',
    'Nechcem odpovedať': 'Nechcem odpovedať'
}

days_map = {
    'Menej ako 1 deň': 'Menej ako 1 deň',
    '1 deň': '1 deň',
    '2 dni': '2 dni',
    '3 dni': '3 dni',
    'Viac ako 3 dni': 'Viac ako 3 dni'
}

reasons_map = {
    'Mala som bolesti': 'Bolesť',
    'Nemala som možnosť sa hygienicky upraviť v škole': 'Nemala som možnosť sa hygienicky upraviť v škole',
    'Nemala som hygienické pomôcky': 'Nemala som hygienické pomôcky',
    'Iné': 'Iný dôvod',
    'Hanbila som sa': 'Hanbila som sa'
}

# --- AFTER 1: Age distribution ---
age_counts = after_data['Vek'].value_counts()

plt.figure(figsize=(8, 5))
bars = plt.bar(age_counts.index, age_counts.values, color=CHART_COLOR)
plt.xlabel('Vek')
plt.title('Rozdelenie veku respondentiek')
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().yaxis.set_visible(False)
total_after = sum(age_counts.values)
for i, v in enumerate(age_counts.values):
    plt.text(i, v + 0.5, f"$\\mathbf{{{v}}}$ {v/total_after*100:.1f}%", ha='center', fontsize=10)
plt.tight_layout()
img_after_age = save_fig('after_age')

# --- AFTER 2: Missed school ---
missed_after = after_data['Chýbala si niekedy v škole kvôli menštruácii?'].map(answer_map_after).value_counts()
missed_after = missed_after.reindex(['Áno', 'Nie', 'Nechcem odpovedať'])

plt.figure(figsize=(8, 5))
bars = plt.barh(missed_after.index, missed_after.values, color=CHART_COLOR)
plt.title('Chýbali ste niekedy v škole kvôli menštruácii?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_ma = sum(missed_after.values)
for i, v in enumerate(missed_after.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_ma*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_missed = save_fig('after_missed_school')

# --- AFTER 3: Days missed ---
days_missed = after_data['Koľko dní si vymeškala počas menštruácii?'].map(days_map).value_counts()
order_days = ['Menej ako 1 deň', '1 deň', '2 dni', '3 dni', 'Viac ako 3 dni']
days_missed = days_missed.reindex([x for x in order_days if x in days_missed.index])

plt.figure(figsize=(8, 5))
bars = plt.barh(days_missed.index, days_missed.values, color=CHART_COLOR)
plt.title('Koľko dní ste chýbali kvôli menštruácii?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_dm = sum(days_missed.values)
for i, v in enumerate(days_missed.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_dm*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_days = save_fig('after_days_missed')

# --- AFTER 4: Reason for absence ---
reasons = after_data['Dôvod tvojej absencie počas menštruácii?'].map(reasons_map).value_counts()

plt.figure(figsize=(8, 5))
bars = plt.barh(reasons.index, reasons.values, color=CHART_COLOR)
plt.title('Dôvod absencie počas menštruácie')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_r = sum(reasons.values)
for i, v in enumerate(reasons.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_r*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_reasons = save_fig('after_reasons')

# --- AFTER 5: Used free pads ---
used_pads = after_data['Používala si bezplatné vložky poskytované v škole?'].map(answer_map_after).value_counts()
used_pads = used_pads.reindex(['Áno', 'Nie', 'Nechcem odpovedať'])

plt.figure(figsize=(8, 5))
bars = plt.barh(used_pads.index, used_pads.values, color=CHART_COLOR)
plt.title('Používali ste bezplatné vložky poskytované v škole?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_up = sum(used_pads.values)
for i, v in enumerate(used_pads.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_up*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_used_pads = save_fig('after_used_pads')

# --- AFTER 6: Products used (detailed) ---
products_map = {
    'Ano, viackrát': 'Áno, viackrát',
    'Ano, raz': 'Áno, raz',
    'Nie': 'Nie',
    'Vedela som o nich, ale nepotrebovala som ich': 'Vedela som o nich, ale nepotrebovala som ich',
    'Nevedela som, že sú dostupné': 'Nevedela som, že sú dostupné'
}
products = after_data['Využili ste niekedy menštruačné pomôcky, ktoré boli v rámci projektu zdarma k dispozícii na škole?'].map(products_map).value_counts()
order_products = ['Áno, viackrát', 'Áno, raz', 'Nie', 'Vedela som o nich, ale nepotrebovala som ich', 'Nevedela som, že sú dostupné']
products = products.reindex([x for x in order_products if x in products.index])

plt.figure(figsize=(10, 5))
bars = plt.barh(products.index, products.values, color=CHART_COLOR)
plt.title('Využili ste niekedy menštruačné pomôcky poskytované v rámci projektu zdarma v škole?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_p = sum(products.values)
for i, v in enumerate(products.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_p*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_products = save_fig('after_products_detail')

# --- AFTER 7: Attendance affected ---
attendance_map = {
    'Ano, chodila som do školy častejšie': 'Áno, chodila som do školy častejšie',
    'Nie, nezmenilo sa to': 'Nie, nezmenilo sa to',
    'Neviem posúdiť': 'Neviem posúdiť'
}
attendance = after_data['Ovplyvnilo to tvoju dochádzku do školy počas menštruácie?'].map(attendance_map).value_counts()
order_att = ['Áno, chodila som do školy častejšie', 'Nie, nezmenilo sa to', 'Neviem posúdiť']
attendance = attendance.reindex([x for x in order_att if x in attendance.index])

plt.figure(figsize=(8, 5))
bars = plt.barh(attendance.index, attendance.values, color=CHART_COLOR)
plt.title('Ovplyvnilo to vašu dochádzku do školy počas menštruácie?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_att = sum(attendance.values)
for i, v in enumerate(attendance.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_att*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_attendance = save_fig('after_attendance')

# --- AFTER 8: Feelings ---
feelings_map = {
    'Lepšie ako predtým': 'Lepšie ako predtým',
    'Rovnako': 'Rovnako',
    'Horšie': 'Horšie'
}
feelings = after_data['Ako sa cítiš počas menštruácie v škole teraz (počas projektu)?'].map(feelings_map).value_counts()
order_f = ['Lepšie ako predtým', 'Rovnako', 'Horšie']
feelings = feelings.reindex([x for x in order_f if x in feelings.index])

plt.figure(figsize=(8, 5))
bars = plt.barh(feelings.index, feelings.values, color=CHART_COLOR)
plt.title('Ako sa cítite počas menštruácie v škole teraz (počas projektu)?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_f = sum(feelings.values)
for i, v in enumerate(feelings.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_f*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_feelings = save_fig('after_feelings')

# --- AFTER 9: Confident ---
confident_map = {
    'Ano': 'Áno',
    'Nie': 'Nie',
    'Neviem': 'Neviem'
}
confident = after_data['Cítiš sa istejšie, keď vieš, že máš v škole k dispozícii hygienické pomôcky?'].map(confident_map).value_counts()
order_c = ['Áno', 'Nie', 'Neviem']
confident = confident.reindex([x for x in order_c if x in confident.index])

plt.figure(figsize=(8, 5))
bars = plt.barh(confident.index, confident.values, color=CHART_COLOR)
plt.title('Cítite sa istejšie, keď viete, že máte v škole k dispozícii hygienické pomôcky?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_c = sum(confident.values)
for i, v in enumerate(confident.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_c*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_confident = save_fig('after_confident')

# --- AFTER 10: Continue project ---
continue_map = {'Ano': 'Áno', 'Je mi to jedno': 'Je mi to jedno'}
continue_proj = after_data['Chcela by si, aby sa poskytovanie vložiek na škole zachovalo aj naďalej?'].map(continue_map).value_counts()
continue_proj = continue_proj.reindex(['Áno', 'Je mi to jedno'])

plt.figure(figsize=(8, 5))
bars = plt.barh(continue_proj.index, continue_proj.values, color=CHART_COLOR)
plt.title('Chceli by ste, aby sa poskytovanie vložiek na škole zachovalo aj naďalej?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_cp = sum(continue_proj.values)
for i, v in enumerate(continue_proj.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_cp*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_continue = save_fig('after_continue')

# --- AFTER 11: Future years ---
future_map = {'Ano, určite': 'Áno, určite', 'Možno': 'Možno'}
future_proj = after_data['Chcela by si, aby boli vložky zadarmo poskytované aj v ďalších školských rokoch?'].map(future_map).value_counts()
future_proj = future_proj.reindex(['Áno, určite', 'Možno'])

plt.figure(figsize=(8, 5))
bars = plt.barh(future_proj.index, future_proj.values, color=CHART_COLOR)
plt.title('Chceli by ste, aby boli vložky zadarmo poskytované aj v ďalších školských rokoch?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_fp = sum(future_proj.values)
for i, v in enumerate(future_proj.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_fp*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_future = save_fig('after_future')

# --- AFTER 12: Discussion ---
discussion_map = {
    'Určite ano': 'Určite áno',
    'Skôr ano': 'Skôr áno',
    'Skôr nie': 'Skôr nie',
    'Určite nie': 'Určite nie'
}
discussion = after_data['Myslíš si, že projekt prispel k tomu, aby sa o menštruácii v škole hovorilo otvorenejšie a prirodzenejšie?'].map(discussion_map).value_counts()
order_d = ['Určite áno', 'Skôr áno', 'Skôr nie', 'Určite nie']
discussion = discussion.reindex([x for x in order_d if x in discussion.index])

plt.figure(figsize=(10, 5))
bars = plt.barh(discussion.index, discussion.values, color=CHART_COLOR)
plt.title('Myslíte si, že projekt prispel k otvorenejšej diskusii o menštruácii v škole?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_d = sum(discussion.values)
for i, v in enumerate(discussion.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_d*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_discussion = save_fig('after_discussion')

# --- AFTER 13: Psych better ---
after_data['Cítila si sa vďaka projektu psychicky lepšie?'] = after_data['Cítila si sa vďaka projektu psychicky lepšie?'].str.capitalize()
psych_map = {'Ano': 'Áno', 'Nie': 'Nie', 'Čiastočne': 'Čiastočne', 'Neviem': 'Neviem'}
psych = after_data['Cítila si sa vďaka projektu psychicky lepšie?'].map(psych_map).value_counts()
order_ps = ['Áno', 'Čiastočne', 'Neviem', 'Nie']
psych = psych.reindex([x for x in order_ps if x in psych.index])

plt.figure(figsize=(8, 5))
bars = plt.barh(psych.index, psych.values, color=CHART_COLOR)
plt.title('Cítili ste sa vďaka projektu psychicky lepšie?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_ps = sum(psych.values)
for i, v in enumerate(psych.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_ps*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_psych = save_fig('after_psych')

# --- AFTER 14: Lectures ---
lectures_map = {
    'Určite ano': 'Určite áno',
    'Skôr ano': 'Skôr áno',
    'Neviem posúdiť': 'Neviem posúdiť',
    'Skôr nie': 'Skôr nie',
    'Určite nie': 'Určite nie'
}
lectures = after_data['V mesiaci december 2025, sa prebehla vo Vašej škola séria prednášok, na tému: Dospievanie, menštruácia a menštruačná chudoba. Prednášali ti: My mami n.o., Zdravé regióny, DM Drogerie a ČLOVEK v ohrození n.o. Pomohli ti tieto aktivity získať nové informácie alebo iný pohľad na túto tému?'].map(lectures_map).value_counts()
order_l = ['Určite áno', 'Skôr áno', 'Neviem posúdiť', 'Skôr nie', 'Určite nie']
lectures = lectures.reindex([x for x in order_l if x in lectures.index])

plt.figure(figsize=(8, 5))
bars = plt.barh(lectures.index, lectures.values, color=CHART_COLOR)
plt.title('Pomohli vám prednášky získať nové informácie alebo iný pohľad na túto tému?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_l = sum(lectures.values)
for i, v in enumerate(lectures.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_l*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_lectures = save_fig('after_lectures')

# --- AFTER 15: Help with issue ---
help_map = {
    'Cítila som sa pokojnejšie a bezpečnejšie': 'Cítila som sa pokojnejšie a bezpečnejšie',
    'Pomohlo mi to vyhnúť sa pretečeniu/nepríjemnosťam': 'Pomohlo mi vyhnúť sa pretečeniu/nepríjemnostiam',
    'Nemala som pri sebe pomôcku a pomohlo mi to prekonať stres': 'Nemala som pri sebe pomôcku, pomohlo mi to prekonať stres',
    'Pomohlo mi to s infekciami alebo zdravotným diskomfortom': 'Pomohlo mi to s infekciami alebo zdravotným diskomfortom',
    'Nepomohlo / nič z toho sa ma netýka': 'Nepomohlo / nič z toho sa ma netýka',
    'Iné': 'Iné'
}
help_issue = after_data['Ak áno, pomohlo ti to vyriešiť niektorý konkrétny problém?'].map(help_map).value_counts()
order_h = ['Cítila som sa pokojnejšie a bezpečnejšie', 'Pomohlo mi vyhnúť sa pretečeniu/nepríjemnostiam',
           'Nemala som pri sebe pomôcku, pomohlo mi to prekonať stres',
           'Pomohlo mi to s infekciami alebo zdravotným diskomfortom', 'Nepomohlo / nič z toho sa ma netýka', 'Iné']
help_issue = help_issue.reindex([x for x in order_h if x in help_issue.index])

plt.figure(figsize=(10, 5))
bars = plt.barh(help_issue.index, help_issue.values, color=CHART_COLOR)
plt.title('Pomohlo vám to vyriešiť niektorý konkrétny problém?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
total_h = sum(help_issue.values)
for i, v in enumerate(help_issue.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{v}}}$ {v/total_h*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_help = save_fig('after_help')

# --- AFTER 16: Future topics ---
topic_columns = {
    'Téme do budúcna: Gynekologické problémy a prevencia': 'Gynekologické problémy a prevencia',
    'Téma do budúcna: Telesné zmeny v období dospievania': 'Telesné zmeny v období dospievania',
    'Téma do budúcna: Vzťah menštruácie a psychického zdravia': 'Vzťah menštruácie a psychického zdravia',
    'Téma do budúcna: Starostlivosť počas menštruácie': 'Starostlivosť počas menštruácie',
    'Téma do budúcnosti: Práva a dôstojnosť žien': 'Práva a dôstojnosť žien',
    'Téma do budúcna: iné': 'Iné'
}
topic_counts = {}
for sk_col, sk_label in topic_columns.items():
    if sk_col in after_data.columns:
        topic_counts[sk_label] = after_data[sk_col].sum()
topics = pd.Series(topic_counts).sort_values(ascending=False)

plt.figure(figsize=(10, 5))
bars = plt.barh(topics.index, topics.values, color=CHART_COLOR)
plt.title('Aké témy by ste do budúcna uvítali na prednáškach?')
plt.gca().invert_yaxis()
for spine in plt.gca().spines.values():
    spine.set_visible(False)
plt.gca().xaxis.set_visible(False)
for i, v in enumerate(topics.values):
    plt.text(v + 0.5, i, f"$\\mathbf{{{int(v)}}}$ {v/num_after*100:.1f}%", va='center', fontsize=10)
plt.tight_layout()
img_after_topics = save_fig('after_topics')


# ═══════════════════════════════════════════
# CROSS-ANALYSIS CHARTS
# ═══════════════════════════════════════════

# Filter pre_data to high school only for comparison
pre_hs = pre_data[pre_data['Škola'] != 'Základnú školu']

yes_no_cross = {
    'Áno': 'Áno', 'Ano': 'Áno',
    'Nie': 'Nie',
    'Niekedy': 'Niekedy',
    'Nechcem odpovedať': 'Nechcem odpovedať'
}

pre_absence = pre_hs['Vynechali ste niekedy školu kvôli menštruácii?'].map(yes_no_cross).value_counts(normalize=True) * 100
post_absence = after_data['Chýbala si niekedy v škole kvôli menštruácii?'].map(yes_no_cross).value_counts(normalize=True) * 100

pre_yes = pre_absence.get('Áno', 0)
post_yes = post_absence.get('Áno', 0)
pre_no = pre_absence.get('Nie', 0)
post_no = post_absence.get('Nie', 0)

# --- CROSS 1: School absence comparison ---
fig, ax = plt.subplots(figsize=(10, 6))
categories = ['Áno', 'Nie']
pre_values = [pre_yes, pre_no]
post_values = [post_yes, post_no]
x = np.arange(len(categories))
width = 0.35
bars1 = ax.bar(x - width/2, pre_values, width, label='Pred inštaláciou', color=COLORS_COMPARISON[0])
bars2 = ax.bar(x + width/2, post_values, width, label='Po inštalácii', color=COLORS_COMPARISON[1])
ax.bar_label(bars1, padding=3, labels=[f'{v:.1f}%' for v in pre_values], fontsize=11, fontweight='bold')
ax.bar_label(bars2, padding=3, labels=[f'{v:.1f}%' for v in post_values], fontsize=11, fontweight='bold')
ax.set_title('Chýbanie v škole kvôli menštruácii', fontsize=14, fontweight='bold')
ax.set_xticks(x)
ax.set_xticklabels(categories)
ax.legend()
change = post_yes - pre_yes
ax.annotate(f'Zmena: {change:+.1f}pb', xy=(0, max(pre_yes, post_yes) + 5), fontsize=12, ha='center',
            color='green' if change < 0 else 'red')
for spine in ax.spines.values():
    spine.set_visible(False)
ax.yaxis.set_visible(False)
plt.tight_layout()
img_cross_absence = save_fig('cross_absence')

# --- CROSS 2: Satisfaction metrics ---
usage_col = 'Využili ste niekedy menštruačné pomôcky, ktoré boli v rámci projektu zdarma k dispozícii na škole?'
usage = after_data[usage_col].value_counts()
used_multiple = usage.get('Ano, viackrát', 0)
used_once = usage.get('Ano, raz', 0)
total_used = used_multiple + used_once

useful = after_data['Mala si pocit, že projekt bol pre dievčatá užitočný?'].value_counts()
useful_yes = useful.get('Ano', 0)

continue_proj_raw = after_data['Chcela by si, aby sa poskytovanie vložiek na škole zachovalo aj naďalej?'].value_counts()
continue_yes_raw = continue_proj_raw.get('Ano', 0)

future_raw = after_data['Chcela by si, aby boli vložky zadarmo poskytované aj v ďalších školských rokoch?'].value_counts()
future_yes_raw = future_raw.get('Ano, určite', 0)
future_maybe_raw = future_raw.get('Možno', 0)

fig, ax = plt.subplots(figsize=(12, 5))
metrics = [
    'Využili bezplatné pomôcky\naspoň raz',
    'Projekt bol užitočný\npre dievčatá',
    'Chcú pokračovanie\nprojektu',
    'Chcú bezplatné pomôcky\naj v ďalších rokoch'
]
values = [
    total_used / num_after * 100,
    useful_yes / num_after * 100,
    continue_yes_raw / num_after * 100,
    (future_yes_raw + future_maybe_raw) / num_after * 100
]
bars = ax.barh(metrics, values, color=CHART_COLOR)
ax.bar_label(bars, padding=3, labels=[f'$\\mathbf{{{v:.1f}}}$%' for v in values])
ax.set_title('Ukazovatele spokojnosti s projektom', fontsize=14, fontweight='bold')
ax.set_xlim(0, 110)
ax.invert_yaxis()
for spine in ax.spines.values():
    spine.set_visible(False)
ax.xaxis.set_visible(False)
plt.tight_layout()
img_cross_satisfaction = save_fig('cross_satisfaction')


# ═══════════════════════════════════════════
# BUILD DOCX
# ═══════════════════════════════════════════

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

style_heading = doc.styles['Heading 1']
style_heading.font.color.rgb = RGBColor(0x1a, 0x4a, 0x6e)

style_heading2 = doc.styles['Heading 2']
style_heading2.font.color.rgb = RGBColor(0x1a, 0x4a, 0x6e)

# --- Helper ---
def add_chart(doc, img_path, width=Inches(6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)

def add_outcome(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(10)


# ═══════════════ TITLE PAGE ═══════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run = p.add_run('OZ Different')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x4a, 0x6e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dátová analýza výskumu menštruačnej chudoby v Bardejove')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()


# ═══════════════ COLLECTED DATA ═══════════════
doc.add_heading('Zozbierané dáta', level=1)

doc.add_heading('Pred inštaláciou menštruačných skriniek:', level=2)
add_bullet(doc, f'{num_pre} respondentiek')
add_bullet(doc, '2 školy (stredná odborná škola + základná škola)')

doc.add_heading('Po inštalácii menštruačných skriniek:', level=2)
add_bullet(doc, f'{num_after} respondentiek')
add_bullet(doc, '1 škola (stredná odborná škola)')

doc.add_page_break()


# ═══════════════ BEFORE INSTALLATION ═══════════════
doc.add_heading('Pred inštaláciou menštruačných skriniek', level=1)

# Age distribution
doc.add_heading('Rozdelenie veku', level=2)
add_chart(doc, img_pre_age)
add_outcome(doc, f'Zo {num_pre} respondentiek bol priemerný vek {avg_age} rokov. Najmladšia respondentka mala 12 rokov, najstaršia 19 rokov. Najväčšie zastúpenie mali 16-ročné respondentky.')

# Age of first period
doc.add_heading('Vek prvej menštruácie', level=2)
add_chart(doc, img_pre_first_period)
add_outcome(doc, f'Priemerný vek prvej menštruácie bol {avg_first_period_age} rokov. Najmladšia respondentka dostala prvú menštruáciu v 9 rokoch, najstaršia v 15 rokoch. Najčastejšie sa prvá menštruácia objavila v 11 a 13 rokoch.')

# Missed school
doc.add_heading('Vynechanie školy kvôli menštruácii', level=2)
add_chart(doc, img_pre_missed)
missed_yes = missed_counts.get('Áno', 0)
add_outcome(doc, f'{missed_yes} respondentiek ({missed_yes/num_pre*100:.1f}%) uviedlo, že niekedy vynechalo školu kvôli menštruácii. Ide o takmer dve tretiny všetkých respondentiek.')

# Affordability
doc.add_heading('Dostupnosť menštruačných pomôcok', level=2)
add_chart(doc, img_pre_afford)
afford_yes_val = afford_counts.get('Áno', 0)
add_outcome(doc, f'{afford_yes_val} respondentiek ({afford_yes_val/num_pre*100:.1f}%) uviedlo, že si aspoň raz nemohli dovoliť kúpiť menštruačné pomôcky z finančných dôvodov.')

# Information preparedness
doc.add_heading('Informovanosť o menštruácii', level=2)
add_chart(doc, img_pre_info_prep)
no_info = info_prep_counts.get('Nemala som žiadne informácie', 0)
partial_info = info_prep_counts.get('Mala som len čiastočné informácie', 0)
add_outcome(doc, f'{no_info} respondentiek ({no_info/num_pre*100:.1f}%) nemalo žiadne informácie pred prvou menštruáciou a {partial_info} ({partial_info/num_pre*100:.1f}%) malo len čiastočné informácie. Spolu viac ako polovica respondentiek nebola dostatočne informovaná.')

# Information sources
doc.add_heading('Zdroje informácií o menštruácii', level=2)
add_chart(doc, img_pre_info_sources)
add_outcome(doc, 'Hlavným zdrojom informácií o menštruácii bola mama (88,0%). Škola (16,5%) a internet (15,8%) boli ďalšími zdrojmi. Prednášky a workshopy boli zdrojom informácií len pre 5,3% respondentiek.')

# Info preparedness vs age hypothesis
doc.add_heading('Informovanosť a vek prvej menštruácie', level=2)
add_chart(doc, img_pre_info_age)
add_outcome(doc, 'Respondentky, ktoré dostali menštruáciu skôr, mali k dispozícii menej informácií. Priemerný vek prvej menštruácie bol 11,7 roka u tých bez informácií, 11,8 roka u čiastočne informovaných a 12,5 roka u plne informovaných.')

# Products used
doc.add_heading('Používané menštruačné pomôcky', level=2)
add_chart(doc, img_pre_products)
add_outcome(doc, 'Menštruačné vložky používalo 97,0% respondentiek. Tampóny používalo 19,5%, intímky a menštruačné nohavičky po 9,0%. Jedna respondentka používala handry.')

# Access to amenities
doc.add_heading('Prístup k vybavenosti', level=2)
add_chart(doc, img_pre_amenities)
add_outcome(doc, f'{full_access} respondentiek ({full_access/num_pre*100:.1f}%) malo plný prístup ku všetkým vybavenostiam. {lacking_any} respondentiek ({lacking_any/num_pre*100:.1f}%) nemalo prístup aspoň k jednej zo základných vybaveností (kúrenie, teplá voda, sprcha/vaňa, splachovací WC).')

# Amenities by siblings
doc.add_heading('Vybavenosť podľa počtu súrodencov', level=2)
add_chart(doc, img_pre_siblings)
add_outcome(doc, 'Bola zistená korelácia 0,4 medzi počtom súrodencov a nedostatkom vybaveností. Respondentky s 5+ súrodencami nemali v priemere 1,25 vybavenosti, zatiaľ čo respondentky bez súrodencov nemali žiadny nedostatok.')

# Amenities by age
doc.add_heading('Vybavenosť podľa veku', level=2)
add_chart(doc, img_pre_age_amenities)
corr_age_lack = pre_data['Vek'].corr(pre_data['Lack_count'])
add_outcome(doc, f'Bola zistená negatívna korelácia -0,39 medzi vekom a nedostatkom vybaveností. Mladšie respondentky (12-13 rokov) mali v priemere 1,33 chýbajúcich vybaveností, zatiaľ čo staršie (18-19 rokov) len 0,03.')

# Symptoms
doc.add_heading('Symptómy počas menštruácie', level=2)
add_chart(doc, img_pre_symptoms)
add_outcome(doc, 'Najčastejším symptómom bol hnev, nervozita, náladovosť a stres (57,9%). Bolesť pociťovalo 30,8%, smútok, depresiu a úzkosť 25,6% a únavu 18,8% respondentiek.')

# Tampon users + hot water
doc.add_heading('Prístup k teplej vode medzi používateľkami tampónov', level=2)
add_chart(doc, img_pre_tampon_water)
tampon_no_water = hot_water_counts.get('Nie', 0)
add_outcome(doc, f'Z {total_tampon} používateliek tampónov {tampon_no_water} ({tampon_no_water/total_tampon*100:.1f}%) nemalo prístup k teplej vode, čo predstavuje hygienické riziko.')

doc.add_page_break()


# ═══════════════ SUMMARY - BEFORE ═══════════════
doc.add_heading('Zhrnutie zistení – pred inštaláciou', level=1)
p = doc.add_paragraph(f'Z {num_pre} respondentiek:')
add_bullet(doc, f'Najmladší vek prvej menštruácie bol 9 rokov')
add_bullet(doc, f'63,2% vynechalo školu kvôli menštruácii')
add_bullet(doc, f'12,0% si nemohlo dovoliť menštruačné pomôcky')
add_bullet(doc, f'26,3% nemalo žiadne informácie pred prvou menštruáciou')
add_bullet(doc, f'97% používa menštruačné vložky')
add_bullet(doc, f'18% má obmedzený prístup k základnej vybavenosti')
add_bullet(doc, f'Mladšie respondentky a respondentky s viac súrodencami majú väčší nedostatok vybaveností')
add_bullet(doc, f'Respondentky s nižším vekom prvej menštruácie mali menej informácií')

doc.add_page_break()


# ═══════════════ AFTER INSTALLATION ═══════════════
doc.add_heading('Po inštalácii menštruačných skriniek', level=1)

# Age
doc.add_heading('Rozdelenie veku', level=2)
add_chart(doc, img_after_age)
add_outcome(doc, f'Z {num_after} respondentiek bolo 66,2% vo veku 16-18 rokov a 33,8% starších ako 18 rokov. 5 respondentiek neuviedlo vek.')

# School absence
doc.add_heading('Absencia v škole', level=2)
add_chart(doc, img_after_missed)
add_chart(doc, img_after_days)
add_chart(doc, img_after_reasons)
add_outcome(doc, '53,2% respondentiek chýbalo v škole kvôli menštruácii. Najčastejšie chýbali 1 deň (42,6%) alebo menej ako 1 deň (31,1%). Dominantným dôvodom bola bolesť (86,9%).')

# Used free pads
doc.add_heading('Používanie bezplatných vložiek v škole', level=2)
add_chart(doc, img_after_used_pads)
add_outcome(doc, '42,3% respondentiek používalo bezplatné vložky poskytované v škole. 55,1% ich nepoužívalo.')

# Products used detail
doc.add_heading('Využitie bezplatných menštruačných pomôcok', level=2)
add_chart(doc, img_after_products)
add_outcome(doc, '30,4% respondentiek využilo bezplatné pomôcky viackrát, 17,7% raz. 26,6% o nich vedelo, ale nepotrebovalo ich. Len 1,3% nevedelo o ich dostupnosti.')

# Attendance
doc.add_heading('Vplyv na dochádzku', level=2)
add_chart(doc, img_after_attendance)
add_outcome(doc, '11,4% respondentiek uviedlo, že vďaka projektu chodili do školy častejšie. Pre väčšinu (64,6%) sa dochádzka nezmenila.')

# Feelings
doc.add_heading('Pocity počas menštruácie v škole', level=2)
add_chart(doc, img_after_feelings)
add_outcome(doc, '17,7% respondentiek sa cítilo lepšie ako predtým. 73,4% sa cítilo rovnako. 8,9% uviedlo zhoršenie.')

# Confident
doc.add_heading('Pocit istoty s dostupnými pomôckami', level=2)
add_chart(doc, img_after_confident)
add_outcome(doc, '79,7% respondentiek sa cítilo istejšie, keď vedeli, že majú v škole k dispozícii hygienické pomôcky.')

# Continue + Future
doc.add_heading('Pokračovanie projektu', level=2)
add_chart(doc, img_after_continue)
add_chart(doc, img_after_future)
add_outcome(doc, '86,1% respondentiek chce, aby sa poskytovanie vložiek zachovalo. 87,3% chce bezplatné pomôcky aj v ďalších školských rokoch. Žiadna respondentka nebola vyslovene proti.')

# Discussion
doc.add_heading('Vplyv na otvorenosť diskusie', level=2)
add_chart(doc, img_after_discussion)
add_outcome(doc, '55,7% respondentiek si myslí, že projekt určite prispel k otvorenejšej diskusii o menštruácii v škole. Spolu so "skôr áno" je to 88,6%.')

# Psych
doc.add_heading('Psychologický prínos projektu', level=2)
add_chart(doc, img_after_psych)
add_outcome(doc, '35,4% respondentiek sa cítilo psychicky lepšie vďaka projektu, 26,6% čiastočne. Spolu 62,0% respondentiek vnímalo pozitívny psychologický vplyv.')

# Lectures
doc.add_heading('Prínos prednášok', level=2)
add_chart(doc, img_after_lectures)
add_outcome(doc, '36,7% respondentiek uviedlo, že prednášky im určite pomohli získať nové informácie. Spolu so "skôr áno" je to 65,8%.')

# Help with issue
doc.add_heading('Riešenie konkrétnych problémov', level=2)
add_chart(doc, img_after_help)
add_outcome(doc, '25,3% respondentiek sa cítilo pokojnejšie a bezpečnejšie. 21,5% sa vyhlo pretečeniu alebo nepríjemnostiam. 11,4% prekonalo stres z nedostatku pomôcok.')

# Future topics
doc.add_heading('Témy pre budúce prednášky', level=2)
add_chart(doc, img_after_topics)
add_outcome(doc, 'Najžiadanejšou témou sú gynekologické problémy a prevencia, nasledované právami a dôstojnosťou žien a starostlivosťou počas menštruácie.')

doc.add_page_break()


# ═══════════════ CROSS ANALYSIS ═══════════════
doc.add_heading('Krížová analýza: Pred vs Po inštalácii', level=1)

# Absence comparison
doc.add_heading('Porovnanie absencie v škole', level=2)
add_chart(doc, img_cross_absence)
add_outcome(doc, f'Absencia v škole kvôli menštruácii klesla z {pre_yes:.1f}% na {post_yes:.1f}%, čo predstavuje pokles o {abs(change):.1f} percentuálnych bodov.')

# Satisfaction
doc.add_heading('Ukazovatele spokojnosti s projektom', level=2)
add_chart(doc, img_cross_satisfaction)
add_outcome(doc, f'48,1% respondentiek využilo bezplatné pomôcky aspoň raz. 88,6% považovalo projekt za užitočný. 86,1% chce pokračovanie projektu a 100% respondentiek chce bezplatné pomôcky aj v budúcich rokoch.')

doc.add_page_break()


# ═══════════════ FINAL SUMMARY ═══════════════
doc.add_heading('Záverečné zhrnutie', level=1)

doc.add_heading('Absencia v škole', level=2)
add_bullet(doc, f'Pred inštaláciou: {pre_yes:.1f}% respondentiek chýbalo v škole kvôli menštruácii')
add_bullet(doc, f'Po inštalácii: {post_yes:.1f}% respondentiek chýbalo v škole kvôli menštruácii')
add_bullet(doc, f'Zmena: pokles o {abs(change):.1f} percentuálnych bodov')

doc.add_heading('Riešenie existujúcich výziev', level=2)
add_bullet(doc, 'Pred: 9,5% si nemohlo dovoliť menštruačné pomôcky')
add_bullet(doc, 'Po: 48,1% využilo bezplatné pomôcky v škole')
add_bullet(doc, 'Po: 79,7% sa cíti istejšie s dostupnými pomôckami')

doc.add_heading('Psychologický dopad', level=2)
add_bullet(doc, 'Pred: 55,8% pociťovalo stres pri výmene pomôcok mimo domova')
add_bullet(doc, 'Po: 62,0% sa cítilo psychicky lepšie vďaka projektu')
add_bullet(doc, 'Po: 25,3% sa cítilo pokojnejšie a bezpečnejšie')

doc.add_heading('Otvorenosť a vzdelávanie', level=2)
add_bullet(doc, 'Pred: 40,0% malo nedostatočné informácie pred prvou menštruáciou')
add_bullet(doc, 'Po: 88,6% uviedlo, že projekt prispel k otvorenejšej diskusii')
add_bullet(doc, 'Po: 65,8% považovalo prednášky za prínosné')

doc.add_heading('Podpora projektu', level=2)
add_bullet(doc, '88,6% považovalo projekt za užitočný pre dievčatá')
add_bullet(doc, '86,1% chce pokračovanie projektu')
add_bullet(doc, '100% chce bezplatné pomôcky aj v ďalších školských rokoch')


# ═══════════════ SAVE ═══════════════
output_path = os.path.join(BASE, '..', 'OZ Different - dátová analýza.docx')
doc.save(output_path)
print(f"\nDOCX saved to: {os.path.abspath(output_path)}")
print("Done!")
